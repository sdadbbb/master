import os
import json
from datetime import datetime
from page.llmClient import LLMClient
from page.llmXlsxManager import LLMXlsxManager
from page.apiTestManager import ApiTestManager
from page.uiTestCaseManager import UITestCaseManager
from log.logger import LoggerUtil

logger = LoggerUtil.get_logger()

# 用例生成系统提示词
CASE_GENERATION_SYSTEM_PROMPT = """你是一名资深测试工程师，擅长根据需求文档进行需求拆分并编写自动化测试用例。
你需要分析用户提供的需求文档，输出严格的 JSON 格式结果，不要输出任何 JSON 以外的内容。

输出 JSON 结构如下：
{
  "requirements_analysis": "需求拆分与分析说明（Markdown格式）",
  "api_cases": [
    {
      "module": "需求模块名",
      "name": "用例名称",
      "description": "用例描述",
      "priority": "高/中/低",
      "request": {
        "method": "GET/POST/PUT/DELETE",
        "url": "完整URL或相对路径",
        "headers": {"Content-Type": "application/json"},
        "data": {}
      },
      "extract": {"变量名": "data.字段路径"},
      "assert": {
        "status_code": 200,
        "response_time": {"less_than": 3000},
        "json": {"success": true}
      }
    }
  ],
  "ui_cases": [
    {
      "module": "需求模块名",
      "name": "用例名称",
      "description": "用例描述",
      "priority": "高/中/低",
      "url": "起始页面URL",
      "steps": [
        {"action": "go_url", "params": {"url": "https://example.com"}},
        {"action": "input_by_placeholder_only", "params": {"placeholder": "请输入用户名", "text": "admin"}},
        {"action": "click_button_by_text", "params": {"button_text": "登录"}},
        {"action": "assert_text_exists", "params": {"text": "欢迎"}}
      ]
    }
  ]
}

UI 步骤可用 action 列表：
go_url, click_element, input_element, input_by_placeholder, input_by_placeholder_only,
click_button_by_text, click_contains_text, find_element_by_text, check_login_result,
wait_element, wait, assert_text_exists, get_element_text

注意：
1. 根据需求合理拆分模块，每个功能点至少覆盖正向和异常场景
2. 接口用例的 assert 必须包含 status_code
3. UI 用例 steps 要具体可执行
4. 如果需求不涉及某类测试，对应数组可为空
"""


class LLMCaseGenerator:
    """大模型测试用例生成器"""

    def __init__(self):
        self.llm = LLMClient()
        self.xlsx_manager = LLMXlsxManager()
        self.api_manager = ApiTestManager()
        self.ui_manager = UITestCaseManager()

    def read_document(self, filepath):
        """读取需求文档内容，支持 txt/md/docx"""
        ext = os.path.splitext(filepath)[1].lower()

        if ext in ('.txt', '.md', '.json', '.yaml', '.yml'):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()

        if ext == '.docx':
            try:
                from docx import Document
                doc = Document(filepath)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        tables_text.append(' | '.join(cells))
                return '\n'.join(paragraphs + tables_text)
            except Exception as e:
                raise ValueError(f'读取 docx 文件失败: {str(e)}')

        raise ValueError(f'不支持的文件格式: {ext}，请上传 txt/md/docx 文件')

    def generate_from_text(self, requirement_text, extra_prompt=''):
        """
        根据需求文本生成测试用例
        :return: (generation_result, xlsx_filename)
        """
        user_prompt = f"""请分析以下需求文档，进行需求拆分并生成接口测试用例和UI测试用例：

---需求文档开始---
{requirement_text}
---需求文档结束---

{extra_prompt}

请严格按照系统提示的 JSON 格式输出。"""

        messages = [
            {'role': 'system', 'content': CASE_GENERATION_SYSTEM_PROMPT},
            {'role': 'user', 'content': user_prompt}
        ]

        response = self.llm.chat(messages)
        result = LLMClient.extract_json(response)

        if 'api_cases' not in result:
            result['api_cases'] = []
        if 'ui_cases' not in result:
            result['ui_cases'] = []
        if 'requirements_analysis' not in result:
            result['requirements_analysis'] = '未提供需求分析'

        _, filename = self.xlsx_manager.save_cases_to_xlsx(result)
        logger.info(f"生成测试用例: API {len(result['api_cases'])} 个, UI {len(result['ui_cases'])} 个, 文件: {filename}")

        return result, filename

    def generate_from_file(self, filepath, extra_prompt=''):
        """从需求文件生成测试用例"""
        content = self.read_document(filepath)
        return self.generate_from_text(content, extra_prompt)

    def import_to_platform(self, generation_result):
        """
        将生成的用例导入平台（写入 tests.json 和 cases.json）
        :return: 导入统计
        """
        api_imported = []
        ui_imported = []

        for case in generation_result.get('api_cases', []):
            test_data = {
                'name': case.get('name', '未命名接口用例'),
                'description': case.get('description', ''),
                'request': case.get('request', {}),
                'extract': case.get('extract', {}),
                'assert': case.get('assert', {})
            }
            saved = self.api_manager.add_test(test_data)
            api_imported.append(saved)

        for case in generation_result.get('ui_cases', []):
            case_data = {
                'name': case.get('name', '未命名UI用例'),
                'description': case.get('description', ''),
                'url': case.get('url', ''),
                'steps': case.get('steps', [])
            }
            saved = self.ui_manager.add_case(case_data)
            ui_imported.append(saved)

        return {
            'api_count': len(api_imported),
            'ui_count': len(ui_imported),
            'api_cases': api_imported,
            'ui_cases': ui_imported
        }

    def import_from_xlsx(self, filename):
        """从 xlsx 文件导入用例到平台"""
        data = self.xlsx_manager.read_cases_from_xlsx(filename)
        if not data:
            raise ValueError('文件不存在或无法读取')
        return self.import_to_platform(data)
